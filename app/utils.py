from langchain_mistralai import ChatMistralAI
import httpx
import pandas as pd
import logging
import time
import struct
import re
from typing import Optional, Union
import os
try:    # for running interface.py
    from model import PCBCharacteristics
    from logger import setup_logger
except ImportError:  # for running main.py
    from .model import PCBCharacteristics
    from .logger import setup_logger

logger = setup_logger(level=logging.INFO)


def _get_file_path(file) -> str:
    """Return path to file for reading. Supports path string or file-like with .name."""
    if isinstance(file, str):
        return file
    if hasattr(file, "name"):
        return file.name
    return str(file)


def extract_word_data(file) -> str:
    """Извлекает текст из документа Word (.docx): параграфы и таблицы."""
    logger.info("Извлечение данных из Word файла.")
    path = _get_file_path(file)
    try:
        from docx import Document
        doc = Document(path)
        parts = []
        for p in doc.paragraphs:
            text = (p.text or "").strip()
            if text:
                parts.append(text)
        for table in doc.tables:
            rows_text = []
            for row in table.rows:
                cells = [(cell.text or "").strip().replace("\n", " ") for cell in row.cells]
                if any(cells):
                    rows_text.append("\t".join(cells))
            if rows_text:
                parts.append("\n".join(rows_text))
        result = "\n\n".join(parts)
        logger.info("Word: извлечено символов %s, слов %s", len(result), len(result.split()))
        return result
    except Exception as e:
        logger.error("Ошибка чтения Word файла: %s", e)
        raise e


def _convert_doc_to_docx(doc_path: str, out_dir: str) -> str:
    """Конвертирует .doc в .docx через LibreOffice или doc2docx. Возвращает путь к .docx."""
    import os
    import subprocess
    import shutil
    base = os.path.splitext(os.path.basename(doc_path))[0]
    docx_path = os.path.join(out_dir, base + ".docx")
    # 1) LibreOffice (если установлен)
    soffice_candidates = [
        os.path.expandvars(r"%ProgramFiles%\LibreOffice\program\soffice.exe"),
        os.path.expandvars(r"%ProgramFiles(x86)%\LibreOffice\program\soffice.exe"),
        "soffice",
        "libreoffice",
    ]
    for soffice in soffice_candidates:
        if soffice in ("soffice", "libreoffice"):
            exe = shutil.which(soffice)
            if not exe:
                continue
            soffice = exe
        elif not os.path.isfile(soffice):
            continue
        try:
            subprocess.run(
                [
                    soffice,
                    "--headless",
                    "--convert-to", "docx",
                    "--outdir", out_dir,
                    doc_path,
                ],
                check=True,
                capture_output=True,
                timeout=60,
                creationflags=subprocess.CREATE_NO_WINDOW if hasattr(subprocess, "CREATE_NO_WINDOW") else 0,
            )
            if os.path.isfile(docx_path):
                return docx_path
        except (subprocess.CalledProcessError, FileNotFoundError, OSError, subprocess.TimeoutExpired) as e:
            logger.debug("LibreOffice не сработал (%s): %s", soffice, e)
            continue
    # 2) doc2docx (требует установленный Word на Windows или LibreOffice)
    try:
        from doc2docx import convert
        convert(doc_path, docx_path)
        if os.path.isfile(docx_path):
            return docx_path
    except Exception as e:
        logger.debug("doc2docx не сработал: %s", e)
    raise RuntimeError(
        "Не удалось конвертировать .doc в .docx. Установите LibreOffice "
        "(https://www.libreoffice.org/) или сохраните документ как .docx в Word."
    )


def _extract_doc_text_ole(doc_path: str) -> str:
    """Извлекает текст из .doc через OLE-парсинг (без MS Word и LibreOffice).
    Реализует разбор piece table из Word Binary File Format (Word 97-2003).
    """
    try:
        import olefile
    except ImportError:
        logger.debug("olefile не установлен, OLE-извлечение недоступно.")
        return ""
    try:
        if not olefile.isOleFile(doc_path):
            return ""
        with olefile.OleFileIO(doc_path) as ole:
            if not ole.exists("WordDocument"):
                return ""
            wd = ole.openstream("WordDocument").read()
            if len(wd) < 0x01AA:
                return ""

            # FIB flags (offset 10): bit 9 (0x0200) = fWhichTblStm (0→0Table, 1→1Table)
            flags = struct.unpack_from("<H", wd, 10)[0]
            use_1table = bool(flags & 0x0200)

            # ccpText: количество символов основного текста (offset 0x004C)
            ccpText = struct.unpack_from("<I", wd, 0x004C)[0]
            if not ccpText:
                return ""

            # Параметры CLX (piece table) в Table-stream
            fcClx  = struct.unpack_from("<I", wd, 0x01A2)[0]
            lcbClx = struct.unpack_from("<I", wd, 0x01A6)[0]

            table_name = "1Table" if use_1table else "0Table"
            if not ole.exists(table_name):
                return ""
            table = ole.openstream(table_name).read()

            clx = table[fcClx: fcClx + lcbClx]

            # Пропускаем Prc-записи (тип 0x01)
            pos = 0
            while pos < len(clx) and clx[pos] == 0x01:
                cb = struct.unpack_from("<H", clx, pos + 1)[0]
                pos += 3 + cb

            # PlcPcd: тип 0x02, затем 4 байта размера, затем данные
            if pos >= len(clx) or clx[pos] != 0x02:
                return ""
            pcdt_size = struct.unpack_from("<I", clx, pos + 1)[0]
            pcdt = clx[pos + 5: pos + 5 + pcdt_size]

            # Количество кусков: (размер - 4) / 12 (n+1 CP по 4 байта + n PCD по 8 байт)
            n_pieces = (len(pcdt) - 4) // 12
            if n_pieces <= 0:
                return ""

            cps = [struct.unpack_from("<I", pcdt, i * 4)[0] for i in range(n_pieces + 1)]
            cp_base = (n_pieces + 1) * 4

            texts = []
            for i in range(n_pieces):
                pcd = pcdt[cp_base + i * 8: cp_base + i * 8 + 8]
                if len(pcd) < 8:
                    break
                fc_raw = struct.unpack_from("<I", pcd, 2)[0]
                is_ansi = bool(fc_raw & 0x40000000)
                fc = fc_raw & 0x3FFFFFFF
                char_count = cps[i + 1] - cps[i]
                if is_ansi:
                    raw = wd[fc >> 1: (fc >> 1) + char_count]
                    piece_text = raw.decode("cp1252", errors="replace")
                else:
                    raw = wd[fc: fc + char_count * 2]
                    piece_text = raw.decode("utf-16-le", errors="replace")
                texts.append(piece_text)

            raw_text = "".join(texts)
            # Убираем управляющие символы, оставляем переносы строк
            raw_text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", " ", raw_text)
            raw_text = raw_text.replace("\r", "\n")
            raw_text = re.sub(r" {3,}", "  ", raw_text)
            raw_text = re.sub(r"\n{3,}", "\n\n", raw_text)
            result = raw_text.strip()
            if result:
                logger.info(".doc OLE: извлечено %d символов.", len(result))
            return result
    except Exception as e:
        logger.debug("OLE-извлечение .doc не удалось: %s", e)
        return ""


def extract_word97_data(file) -> str:
    """Извлекает текст из Word 97–2003 (.doc).
    Порядок: Win32 COM → LibreOffice/doc2docx → OLE-парсинг (pure Python).
    """
    import os
    import sys
    path = _get_file_path(file)
    path = os.path.abspath(path)
    logger.info("Извлечение данных из Word 97-2003 (.doc).")

    # 1) Win32 COM через установленный MS Word
    if sys.platform == "win32":
        try:
            import win32com.client
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            try:
                doc = word.Documents.Open(path, ReadOnly=True)
                try:
                    text = doc.Content.Text
                    if text and text.strip():
                        logger.info(".doc: текст извлечён через MS Word (COM).")
                        return text.strip()
                finally:
                    doc.Close(False)
            finally:
                word.Quit()
        except Exception as e:
            logger.debug("Извлечение .doc через Word COM не удалось: %s", e)

    # 2) LibreOffice / doc2docx — конвертация в .docx
    import shutil
    import tempfile
    try:
        tmpdir = tempfile.mkdtemp()
        try:
            doc_copy = os.path.join(tmpdir, "input.doc")
            shutil.copy2(path, doc_copy)
            docx_path = _convert_doc_to_docx(doc_copy, tmpdir)
            return extract_word_data(docx_path)
        finally:
            shutil.rmtree(tmpdir, ignore_errors=True)
    except Exception as e:
        logger.debug("Конвертация .doc не удалась: %s", e)

    # 3) Pure-Python OLE-парсинг (без внешних инструментов)
    text = _extract_doc_text_ole(path)
    if text:
        logger.info(".doc: текст извлечён через OLE-парсинг.")
        return text

    raise RuntimeError(
        "Не удалось прочитать .doc файл. Пожалуйста, пересохраните документ как .docx в Microsoft Word."
    )


def extract_document_data(file) -> str:
    """Извлекает текст из Excel, Word или txt по расширению файла."""
    path = _get_file_path(file)
    path_lower = path.lower()
    if path_lower.endswith(".docx"):
        text = extract_word_data(file)
    elif path_lower.endswith(".doc"):
        text = extract_word97_data(file)
    elif path_lower.endswith((".xlsx", ".xlsm", ".xls")):
        text = extract_excel_data(file)
    elif path_lower.endswith(".txt"):
        text = extract_text_data(file)
    else:
        raise ValueError(
            "Неподдерживаемый формат. Используйте файл .xlsx, .xls, .docx, .doc или .txt "
            "(Лист технических требований ПП / бланк заказа)."
        )
    return _truncate_for_llm(text)


# ═══════════════════════════════════════════════════════════════════
# ИЗВЛЕЧЕНИЕ EXCEL
# Бланки заказов — это формы: значения лежат рядом с метками, а скрытые
# листы/строки/колонки содержат инструкции по заполнению и списки
# вариантов выпадающих меню. Всё скрытое в текст для LLM не попадает.
# ═══════════════════════════════════════════════════════════════════

# Максимальный размер текста, передаваемого в LLM (символов).
MAX_EXTRACT_CHARS = int(os.getenv("MAX_EXTRACT_CHARS", "60000") or "60000")

# Листы-шаблоны/инструкции: не содержат данных конкретного заказа,
# зато содержат значения-примеры, которые LLM принимает за реальные.
_NOISE_SHEET_MARKERS = (
    "пример", "example", "sample", "шаблон", "template",
    "инструкция", "instruction", "политика", "policy", "согласие",
)


def _truncate_for_llm(text: str) -> str:
    if len(text) <= MAX_EXTRACT_CHARS:
        return text
    logger.warning(
        "Извлечённый текст слишком большой (%d символов), обрезан до %d.",
        len(text), MAX_EXTRACT_CHARS,
    )
    return text[:MAX_EXTRACT_CHARS] + "\n[... текст обрезан: документ слишком большой ...]"


def _is_noise_sheet(name: str) -> bool:
    low = (name or "").casefold()
    return any(marker in low for marker in _NOISE_SHEET_MARKERS)


def _format_cell_value(v) -> str:
    """Значение ячейки как компактный текст: без хвоста .0, даты как YYYY-MM-DD."""
    import datetime as _dt
    if v is None:
        return ""
    if isinstance(v, bool):
        return "Да" if v else "Нет"
    if isinstance(v, float):
        if v.is_integer():
            return str(int(v))
        return str(round(v, 6))
    if isinstance(v, _dt.datetime):
        if v.hour == v.minute == v.second == 0:
            return v.strftime("%Y-%m-%d")
        return v.strftime("%Y-%m-%d %H:%M")
    if isinstance(v, _dt.date):
        return v.strftime("%Y-%m-%d")
    return re.sub(r"\s+", " ", str(v)).strip()


def _rows_to_text(sheet_title: str, rows) -> str:
    """Рендерит строки листа как 'значение | значение': метка и выбранное значение
    оказываются на одной строке, пустые ячейки не создают шума."""
    lines = []
    for cells in rows:
        vals = [c for c in cells if c]
        if vals:
            lines.append(" | ".join(vals))
    if not lines:
        return ""
    return f"=== Лист: {sheet_title} ===\n" + "\n".join(lines)


def _dv_source_cells(ws) -> set:
    """Ячейки-источники выпадающих списков (data validation type=list) на этом листе.

    В бланках заказа (напр. Резонит) варианты меню лежат в видимых дальних
    колонках; без исключения LLM видит и выбранное значение, и все альтернативы.
    """
    from openpyxl.utils import range_boundaries
    excluded = set()
    try:
        dvs = list(ws.data_validations.dataValidation)
    except Exception:
        return excluded
    for dv in dvs:
        if dv.type != "list" or not dv.formula1:
            continue
        ref = str(dv.formula1).strip().lstrip("=")
        # Только простые ссылки на диапазон ТЕКУЩЕГО листа: $AI$15:$AN$15.
        # Инлайн-списки ("да,нет"), именованные диапазоны и другие листы пропускаем.
        if not re.fullmatch(r"\$?[A-Za-z]{1,3}\$?\d+(?::\$?[A-Za-z]{1,3}\$?\d+)?", ref):
            continue
        try:
            min_c, min_r, max_c, max_r = range_boundaries(ref.replace("$", ""))
        except Exception:
            continue
        if (max_r - min_r + 1) * (max_c - min_c + 1) > 500:
            continue  # защита от ссылок на целые колонки
        for r in range(min_r, max_r + 1):
            for c in range(min_c, max_c + 1):
                excluded.add((r, c))
    return excluded


def _extract_xlsx_structured(path: str) -> str:
    """Извлечение .xlsx через openpyxl с учётом скрытых областей и dropdown-источников."""
    import openpyxl
    from openpyxl.utils import column_index_from_string

    wb = openpyxl.load_workbook(path, data_only=True)
    parts = []
    for ws in wb.worksheets:
        if ws.sheet_state != "visible":
            logger.info("Excel: скрытый лист '%s' пропущен.", ws.title)
            continue
        if _is_noise_sheet(ws.title):
            logger.info("Excel: лист-шаблон/инструкция '%s' пропущен.", ws.title)
            continue

        hidden_cols = set()
        for letter, dim in ws.column_dimensions.items():
            if not dim.hidden:
                continue
            lo = dim.min or column_index_from_string(letter)
            hi = dim.max or lo
            hidden_cols.update(range(lo, hi + 1))
        hidden_rows = {r for r, dim in ws.row_dimensions.items() if dim.hidden}
        dv_cells = _dv_source_cells(ws)

        rows_out = []
        for row in ws.iter_rows():
            if not row or row[0].row in hidden_rows:
                continue
            cells = []
            for cell in row:
                if cell.column in hidden_cols or (cell.row, cell.column) in dv_cells:
                    continue
                cells.append(_format_cell_value(cell.value))
            rows_out.append(cells)

        text = _rows_to_text(ws.title, rows_out)
        if text:
            parts.append(text)
    return "\n\n".join(parts)


def _extract_xls_structured(path: str) -> str:
    """Извлечение .xls через xlrd с учётом скрытых листов/строк/колонок."""
    import xlrd

    def format_xls_cell(book, cell) -> str:
        if cell.ctype == xlrd.XL_CELL_DATE:
            try:
                dt = xlrd.xldate_as_datetime(cell.value, book.datemode)
                return _format_cell_value(dt)
            except Exception:
                return _format_cell_value(cell.value)
        if cell.ctype in (xlrd.XL_CELL_EMPTY, xlrd.XL_CELL_BLANK):
            return ""
        if cell.ctype == xlrd.XL_CELL_BOOLEAN:
            return "Да" if cell.value else "Нет"
        return _format_cell_value(cell.value)

    book = xlrd.open_workbook(path, formatting_info=True)
    parts = []
    for ws in book.sheets():
        if getattr(ws, "visibility", 0) != 0:
            logger.info("Excel: скрытый лист '%s' пропущен.", ws.name)
            continue
        if _is_noise_sheet(ws.name):
            logger.info("Excel: лист-шаблон/инструкция '%s' пропущен.", ws.name)
            continue
        hidden_cols = {c for c, ci in ws.colinfo_map.items() if ci.hidden}
        hidden_rows = {r for r, ri in ws.rowinfo_map.items() if ri.hidden}

        rows_out = []
        for r in range(ws.nrows):
            if r in hidden_rows:
                continue
            cells = []
            for c in range(ws.ncols):
                if c in hidden_cols:
                    continue
                cells.append(format_xls_cell(book, ws.cell(r, c)))
            rows_out.append(cells)

        text = _rows_to_text(ws.name, rows_out)
        if text:
            parts.append(text)
    return "\n\n".join(parts)


def _extract_excel_via_pandas(file) -> str:
    """Fallback-извлечение через pandas (header=None, без скрытых областей)."""
    excel_file = pd.ExcelFile(file)
    all_data = []
    for sheet_name in excel_file.sheet_names:
        if _is_noise_sheet(sheet_name):
            continue
        df = pd.read_excel(excel_file, sheet_name=sheet_name, header=None)
        df = df.dropna(how="all").dropna(axis=1, how="all")
        if df.empty:
            continue
        rows = []
        for _, row in df.iterrows():
            vals = [_format_cell_value(v) for v in row.tolist() if not pd.isna(v)]
            vals = [v for v in vals if v]
            if vals:
                rows.append(" | ".join(vals))
        if rows:
            all_data.append(f"=== Лист: {sheet_name} ===\n" + "\n".join(rows))
    return "\n\n".join(all_data)


def extract_excel_data(file) -> str:
    """Извлекает данные из Excel (.xlsx/.xls) в компактный текст для LLM.

    Скрытые листы, строки и колонки пропускаются (там инструкции по заполнению
    и списки вариантов), источники выпадающих списков исключаются, значения
    рендерятся построчно: 'Метка | значение'.
    """
    logger.info("Starting to extract data from the Excel file.")
    path = _get_file_path(file)
    try:
        if path.lower().endswith((".xlsx", ".xlsm")):
            excel_txt = _extract_xlsx_structured(path)
        else:
            excel_txt = _extract_xls_structured(path)
        if not excel_txt.strip():
            raise ValueError("структурное извлечение вернуло пустой текст")
    except Exception as e:
        logger.warning("Структурное извлечение Excel не удалось (%s), fallback на pandas.", e)
        try:
            excel_txt = _extract_excel_via_pandas(file)
        except Exception as e2:
            logger.error("Error reading Excel file: %s", e2)
            raise e2
    logger.info(
        "Extracted text length: %s, Word count: %s", len(excel_txt), len(excel_txt.split())
    )
    return excel_txt


def extract_text_data(file) -> str:
    """Извлекает текст из .txt, подбирая кодировку (utf-8 → cp1251 → cp866)."""
    path = _get_file_path(file)
    with open(path, "rb") as f:
        raw = f.read()
    for enc in ("utf-8-sig", "utf-8", "cp1251", "cp866"):
        try:
            text = raw.decode(enc)
            break
        except UnicodeDecodeError:
            continue
    else:
        text = raw.decode("utf-8", errors="replace")
    text = text.strip()
    logger.info("TXT: извлечено символов %s.", len(text))
    return text


# ═══════════════════════════════════════════════════════════════════
# СЕТЕВОЙ РЕЖИМ ДЛЯ MISTRAL API (работа с VPN/proxy и без него)
#
# api.mistral.ai в ряде сетей доступен только через VPN/proxy, а иногда —
# наоборот, только напрямую (VPN выключен, но переменные HTTP_PROXY остались
# в окружении и указывают на неработающий локальный порт).
# Режим определяется автоматически один раз и кэшируется; при сетевой ошибке
# кэш сбрасывается, и следующая попытка выбирает режим заново — поэтому
# включение/выключение VPN на ходу не требует перезапуска приложения.
# ═══════════════════════════════════════════════════════════════════

_mistral_trust_env_cache: Optional[bool] = None


def get_mistral_base_url() -> str:
    return (os.getenv("MISTRAL_BASE_URL") or "").strip() or "https://api.mistral.ai/v1"


def reset_mistral_network_mode() -> None:
    """Сбрасывает выбранный сетевой режим (вызывается после сетевой ошибки)."""
    global _mistral_trust_env_cache
    if _mistral_trust_env_cache is not None:
        logger.info("Сброшен кэш сетевого режима Mistral — режим будет выбран заново.")
    _mistral_trust_env_cache = None


def _probe_mistral(trust_env: bool, base_url: str, api_key: str) -> bool:
    """Быстрая проверка доступности Mistral API в заданном сетевом режиме.

    Проверяется именно СЕТЕВАЯ достижимость: ответ 401/403 (неверный ключ)
    тоже означает, что канал до API есть.
    """
    api_key = (api_key or "").strip()
    # Пустой заголовок Authorization httpx отвергает локально (LocalProtocolError),
    # поэтому при отсутствии ключа шлём запрос без него.
    headers = {"Authorization": f"Bearer {api_key}"} if api_key else {}
    try:
        with httpx.Client(
            base_url=base_url,
            timeout=httpx.Timeout(connect=3.0, read=5.0, write=3.0, pool=3.0),
            trust_env=trust_env,
            headers=headers,
        ) as client:
            return client.get("/models").status_code < 500
    except Exception as e:
        logger.debug("Проба Mistral (proxy=%s) не удалась: %s", trust_env, e)
        return False


def resolve_mistral_trust_env(api_key: str = "") -> bool:
    """Определяет, использовать ли системный proxy для запросов к Mistral.

    MISTRAL_TRUST_ENV=1/0 — принудительный режим без автоопределения.
    Иначе: если системный proxy задан, проверяем оба режима и выбираем рабочий.
    """
    global _mistral_trust_env_cache

    forced = (os.getenv("MISTRAL_TRUST_ENV") or "").strip()
    if forced:
        return forced != "0"

    if _mistral_trust_env_cache is not None:
        return _mistral_trust_env_cache

    proxy = (os.getenv("HTTPS_PROXY") or os.getenv("HTTP_PROXY") or "").strip()
    if not proxy:
        _mistral_trust_env_cache = True  # proxy не задан — режимы эквивалентны
        return True

    base_url = get_mistral_base_url()
    if _probe_mistral(True, base_url, api_key):
        logger.info("Mistral API доступен через системный proxy %s — используем его.", proxy)
        _mistral_trust_env_cache = True
    elif _probe_mistral(False, base_url, api_key):
        logger.warning(
            "Mistral API недоступен через proxy %s, но отвечает напрямую — "
            "запросы пойдут в обход proxy (VPN выключен?).", proxy
        )
        _mistral_trust_env_cache = False
    else:
        logger.warning(
            "Mistral API не отвечает ни через proxy %s, ни напрямую. "
            "Проверьте VPN/интернет. Используем proxy-режим.", proxy
        )
        _mistral_trust_env_cache = True
    return _mistral_trust_env_cache


def create_mistral_http_client(api_key: str, read_timeout: float = 120.0) -> httpx.Client:
    """Единая точка создания HTTP-клиента для Mistral API.

    Оба LLM-шага (Промпт 1 в create_pcb_model и Промпт 2 в нормализаторе) обязаны
    использовать одинаковые настройки сети. Иначе один шаг ходит через системный
    proxy, а другой — напрямую, и второй виснет в ConnectTimeout (WinError 10060).

    Важно: если передать client в ChatMistralAI, он НЕ выставляет base_url сам,
    поэтому у клиента должен быть base_url.
    """
    trust_env = resolve_mistral_trust_env(api_key)
    return httpx.Client(
        base_url=get_mistral_base_url(),
        # read увеличен: ответы LLM для больших документов занимают > 30 сек.
        timeout=httpx.Timeout(connect=15.0, read=read_timeout, write=30.0, pool=15.0),
        trust_env=trust_env,
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
            "Accept": "application/json",
        },
    )


def create_pcb_model(params: dict[str, str]) -> ChatMistralAI:
    """Creates and configures a ChatMistralAI model instance for PCB characteristics parsing.

    Args:
        params (dict[str, str]): A dictionary containing parameters for model configuration.
            Expected keys:
                - 'api_key': The API key for authenticating with the ChatMistralAI service.

    Returns:
        ChatMistralAI: An instance of the ChatMistralAI model configured for PCB characteristics parsing.
    """
    api_key = (params.get("api_key") or "").strip()
    # Логируем префикс ключа (безопасно) чтобы понять, какой ключ реально используется в сервере.
    key_prefix = (api_key[:6] + "...") if api_key else None
    logger.info("MISTRAL_API_KEY prefix: %s (len=%s)", key_prefix, len(api_key))
    if api_key == "mistral_api_key":
        logger.warning("Используется плейсхолдер MISTRAL_API_KEY. Проверьте загрузку .env.")
    if not api_key:
        raise ValueError(
            "Mistral API key is empty. Set environment variable `MISTRAL_API_KEY` "
            "before starting the app."
        )

    http_client = create_mistral_http_client(api_key)
    llm = ChatMistralAI(
        model="mistral-medium-latest",
        temperature=0.1,
        api_key=api_key,
        client=http_client,
    )
    return llm.with_structured_output(PCBCharacteristics)


def process_excel_pcb_with_retry(
    excel_txt: str,
    llm_parser: ChatMistralAI,
    max_retries: int = 3,
    delay: float = 2.0,
    model_params: Optional[dict] = None,
) -> Optional[dict]:
    """Processes Excel data for PCB characteristics using a ChatMistralAI model with retry logic.

    Args:
        excel_txt (str): A string containing the Excel data to be processed.
        llm_parser (ChatMistralAI): An instance of the ChatMistralAI model used for parsing PCB data.
        max_retries (int): Maximum number of retry attempts.
        delay (float): Delay between retries in seconds.
        model_params (dict | None): параметры модели ('api_key'). Если переданы, после
            сетевой ошибки клиент пересоздаётся с заново выбранным сетевым режимом —
            это позволяет подхватить включение/выключение VPN без перезапуска сервера.

    Returns:
        dict: A dictionary containing the processed PCB characteristics, or None if all retries failed.
    """
    messages = [
        (
            "system",
            "You are an experienced PCB engineer. Extract PCB characteristics from the provided data. "
            "Data may come from Excel or from a Word 'Technical Requirements Sheet PCB' / 'Лист технических требований ПП' (Russian/English), "
            "or from a Russian order form ('Бланк заказа', 'Технические требования'). "
            "Excel forms are rendered as lines 'Метка | значение [| значение2]' — the chosen value(s) follow their label on the same line. "
            "Extract at least the following fields into the structured PCBCharacteristics object: "
            "company_name (customer: Заказчик / Название фирмы заказчика / Наименование организации), "
            "board_name (board name: Identification / Обозначение / Наименование платы / Название проекта / Название файла платы), "
            "quantity (boards quantity: 'Количество плат в заказе, шт' / 'Требуется изготовить, шт' / 'Количество, шт' / Qty), "
            "base_material (Base material / Тип материала / Материал основания / Базовый материал; keep verbatim incl. Tg, e.g. 'FR4, Tg≥150', 'FR4 HiTg (Tg≥170°C)', 'FR4 IT180A', 'FR4 типовой'), "
            "board_thickness (finished PCB thickness with tolerance: 'Finished thickness with tolerance, mm' / 'Толщина с допуском, мм' / 'Толщина платы, мм' / 'Толщина материала/общая толщина платы'), "
            "foil_thickness (copper foil + plating: 'Толщина фольги, мкм' / 'Толщина меди+покрытие' / 'Thickness CU'; keep verbatim, e.g. '18', '35', '0,035+0,025', '18/35/35/18'), "
            "layer_count ('Количество проводящих слоев' or count conductive layers in the layer stack: Top/Bottom/L1..Ln/Inner/Signal; solder mask, paste, silkscreen/overlay are NOT conductive layers), "
            "board_size (size of a SINGLE board: 'Size with tolerance' / '(Длина x ширина) с допуском' / 'Размер платы (длина/ширина)' / 'Длина платы' + 'Ширина платы'; "
            "output ONLY the two dimensions as 'LENGTH x WIDTH' and DROP every tolerance — "
            "ISO grades (h12, H7, js13), ±0.2, +0,3/-0,3, ±10%. Example: rows 'длина, мм | 114 | Допуск | h12' + 'ширина, мм | 47 | h12' → '114 x 47'), "
            "panelization (free-text summary of panel/заготовка data: gaps, matrix, tech fields), "
            "panel_size (ONLY panel dimensions as 'LENGTH x WIDTH', from 'Размер панели, мм' / 'Габариты панели'; tolerances dropped; NEVER the single-board size), "
            "boards_per_panel (number of boards in one panel: 'Количество ПП в панели' / 'Количество плат в панели'; for an NxM matrix give the product), "
            "different_boards_per_panel (how many DIFFERENT board types the panel holds; a panel repeating one board → '1'), "
            "technological_fields (Yes/No: 'Наличие технологических полей' / 'Размер полей' / 'ширина полей'; also Yes when the panel is larger than the board because of edge rails), "
            "impedance_control (Yes/No: 'Контроль импедансов' / 'Контроль волнового сопротивления' / 'impedance control'; 'есть'/'требуется' → Yes; a stackup table with impedance values in Ohm also → Yes), "
            "min_hole_size (minimum PLATED hole diameter in mm: 'Мин. диаметр металлизир. отверстия' / 'Минимальное металлизированное отверстие' / 'Smallest plated hole size'; number only, e.g. '0.2'), "
            "marking_side (side of the manufacturer marking: 'сторона маркировки' / 'Маркировка изготовителя'; TOP / BOTTOM / TOP+BOTTOM / None; "
            "append the application method when stated, e.g. 'TOP, шелкография'), "
            "date_code (Yes/No: manufacturing date marking required — 'месяц и год изготовления', 'дату изготовления методом шелкографии', 'Manufacturing data'; "
            "a required format such as 'требуется в формате ХХ (римские) ХХ (арабские)' → Yes; append the method when stated, e.g. 'Yes, шелкография'), "
            "serial_number (Yes/No: individual serial or sequential number required per board — 'порядковый номер в партии', 'заводской порядковый номер', 'уникальный код', 'лазерный штрих-код'; a required number format like 'требуется в формате ХХХ' → Yes), "
            "coverage_type (surface finish: 'Surface Finish' / 'Final Finish' / 'Финишное покрытие платы/площадок'; keep verbatim, e.g. 'ENIG', 'Гор.Пос.', 'ПОС-63HAL + Carbon', 'ImmSn', 'Иммерсионное золото (Хим.Н5 Зл0,1)'), "
            "solder_mask_colour (mask presence/sides + colour: 'Наличие маски/цвет' / 'Наличие паяльной маски' / 'Цвет паяльной маски' / solder mask layers colour), "
            "solder_mark_colour (legend presence/sides + colour: 'Наличие маркировки краской/цвет' / 'Маркировка' + 'Цвет маркировки' / silkscreen colour), "
            "electrical_testing ('Электроконтроль' / 'Электротестирование' / 'Провести электроконтроль' / electrical test), "
            "ipc_class (ONLY explicit IPC mentions: 'IPC-A-600 Class 2/3', 'Соответствие: IPC...'; NOTE: 'Класс точности по ГОСТ Р 53429-2009' is a GOST accuracy class, NOT an IPC class — never derive ipc_class from it), "
            "pcb_type ('Тип платы': Жесткая = Rigid, Гибкая = Flex, Гибко-жесткая = Rigid-Flex), "
            "edge_plating ('Торцевая металлизация' / 'Металлизированный торец платы'), "
            "contour_treatment ('Мех обработка контура' / 'Механообработка контура' / 'Метод обработки контура': фрезерование / скрайбирование / V-scoring) "
            "and any other PCB fields defined in the schema. "
            "Ignore form-filling instructions and hints; if the text still contains lists of selectable menu options, extract only the actually chosen value (the one right after the label). "
            "Ignore gerber/drill file name listings except for reading the layer stack. "
            "If any information is missing, use empty string or 0 as appropriate."
        ),
        ("human", excel_txt),
    ]
    
    for attempt in range(max_retries):
        try:
            logger.info("Attempting to process PCB data (attempt %d/%d)", attempt + 1, max_retries)
            answer = llm_parser.invoke(messages)
            logger.info("Successfully processed PCB data")
            return answer.model_dump()

        except Exception as e:
            error_msg = str(e)
            logger.warning("Attempt %d/%d failed: %s", attempt + 1, max_retries, error_msg)

            # Определяем тип ошибки
            is_rate_limit = "429" in error_msg or "capacity exceeded" in error_msg.lower()
            # WinError 10054 = connection reset, 10061 = connection refused — временные сетевые сбои
            is_network_error = any(x in error_msg for x in (
                "10054", "10061", "ReadError", "ConnectError",
                "TimeoutException", "RemoteProtocolError", "ConnectionReset",
                "Connection reset", "forcibly closed",
            ))

            if is_rate_limit or is_network_error:
                if attempt < max_retries - 1:
                    wait_time = delay * (2 ** attempt)  # Exponential backoff
                    reason = "Rate limit" if is_rate_limit else "Network error"
                    logger.info("%s. Retry in %.1f sec...", reason, wait_time)
                    # Сетевая ошибка может означать, что VPN включили/выключили:
                    # сбрасываем выбранный режим и пересобираем клиент.
                    if is_network_error:
                        reset_mistral_network_mode()
                        if model_params:
                            try:
                                llm_parser = create_pcb_model(model_params)
                                logger.info("Клиент Mistral пересоздан с новым сетевым режимом.")
                            except Exception as rebuild_err:
                                logger.warning("Не удалось пересоздать клиент: %s", rebuild_err)
                    time.sleep(wait_time)
                    continue
                else:
                    logger.error("All %d attempts failed.", max_retries)
                    if is_rate_limit:
                        raise Exception("Сервис временно недоступен (превышен лимит запросов). Попробуйте позже.")
                    raise Exception(
                        f"Ошибка сети при обращении к Mistral API: {error_msg}\n"
                        "Проверьте интернет-соединение и настройки proxy/антивируса."
                    )
            else:
                # Нереентерабельная ошибка (неверный ключ, невалидный запрос и т.д.)
                logger.error("Non-retryable error: %s", error_msg)
                raise e

    raise RuntimeError("Не удалось обработать данные PCB после всех попыток.")


